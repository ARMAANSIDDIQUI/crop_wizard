import os
import requests
import urllib.parse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_flowchart(output_path):
    dot_graph = """
    digraph G {
        rankdir=TB;
        fontname="Helvetica,Arial,sans-serif";
        node [fontname="Helvetica,Arial,sans-serif", shape=box, style="filled,rounded", color="#E1E5F2", fontcolor="#1D3557"];
        edge [fontname="Helvetica,Arial,sans-serif", color="#457B9D"];

        A [label="Input Features\n(N, P, K, pH, Rainfall, Temp, Humidity, Soil, Month)", fillcolor="#A8DADC"];
        B [label="Data Preprocessing\n(Label Encoding & Imputation)", fillcolor="#F1FAEE"];
        
        C [label="Voting Classifier (Ensemble Council)", shape=folder, fillcolor="#FFD166"];
        
        D1 [label="CatBoost\n(Gradient Boosting for Categorical Data)", fillcolor="#EF476F", fontcolor="white"];
        D2 [label="XGBoost\n(Gradient Boosting for Numerical Patterns)", fillcolor="#EF476F", fontcolor="white"];
        D3 [label="Random Forest\n(Bagging to Reduce Variance)", fillcolor="#EF476F", fontcolor="white"];
        D4 [label="Ridge Regression\n(L2 Regularization)", fillcolor="#EF476F", fontcolor="white"];
        D5 [label="Lasso Regression\n(L1 Regularization)", fillcolor="#EF476F", fontcolor="white"];
        
        E [label="Majority Vote / Probabilistic Aggregation", fillcolor="#118AB2", fontcolor="white"];
        F [label="Recommended Crop", shape=ellipse, fillcolor="#06D6A0", fontcolor="white", penwidth=2];
        
        A -> B;
        B -> C;
        C -> D1;
        C -> D2;
        C -> D3;
        C -> D4;
        C -> D5;
        
        D1 -> E;
        D2 -> E;
        D3 -> E;
        D4 -> E;
        D5 -> E;
        
        E -> F;
    }
    """
    
    url = f"https://quickchart.io/graphviz?format=png&graph={urllib.parse.quote(dot_graph)}"
    response = requests.get(url)
    if response.status_code == 200:
        with open(output_path, "wb") as f:
            f.write(response.content)
        print(f"Flowchart successfully saved to {output_path}")
    else:
        print(f"Failed to generate flowchart: {response.text}")

def create_word_doc(output_path, image_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('ML Service: Crop Recommendation Algorithm Approach', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "The automated Crop Recommendation system leverages a state-of-the-art ensemble machine learning approach "
        "designed for the highest accuracy and robustness against real-world noise. Instead of relying on a single algorithm, "
        "it creates a 'Council of Experts' (a Voting Classifier) to recommend the optimal crop based on input features "
        "including soil attributes (N, P, K, pH) and environmental conditions (Rainfall, Temperature, Humidity, Soil Type, and planting Month)."
    )
    
    doc.add_heading('2. Algorithm Flowchart', level=1)
    doc.add_paragraph("The visual representation of the prediction pipeline is shown below:")
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Flowchart Image not found]")

    doc.add_heading('3. Data Strategy', level=1)
    p = doc.add_paragraph()
    p.add_run('Realistic Noise Injection: ').bold = True
    p.add_run(
        "To ensure the model performs well in real-world scenarios where sensor data isn't perfect, "
        "the training dataset incorporates controlled Gaussian noise. This prevents the model from achieving unrealistic 100% "
        "training accuracy (overfitting) by forcing it to learn subtle, overlapping conditions rather than memorizing exact ranges."
    )
    
    doc.add_heading('4. Model Architecture: Council of Experts', level=1)
    doc.add_paragraph(
        "The model uses a Voting Classifier grouping five distinct algorithms to ensure robust, non-biased predictions:"
    )
    
    models = [
        ("CatBoost:", " Exceptional for handling categorical variables like soil type efficiently."),
        ("XGBoost:", " A highly optimized Gradient Boosting algorithm specialized in capturing complex, non-linear numerical patterns."),
        ("Random Forest:", " Utilizes bagging to build hundreds of decision trees, reducing variance and improving stability."),
        ("Ridge Regression (L2):", " A regularized linear model preventing any single feature from disproportionately dominating the logic."),
        ("Lasso Regression (L1):", " Another regularization approach that simplifies decision-making by implicitly ignoring irrelevant features.")
    ]
    
    for name, desc in models:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(name).bold = True
        p.add_run(desc)

    doc.add_heading('5. Model Validation', level=1)
    doc.add_paragraph(
        "The model is rigorously validated using 3-Fold Cross-Validation, preventing data memorization and verifying roughly 95% "
        "accuracy on unseen validation subsets. This ensures steady, reliable backend recommendations for the Crop Wizard platform."
    )
    
    doc.add_heading('6. Evaluation Metrics', level=1)
    doc.add_paragraph(
        "Below are the visual representations of the model's performance on the validation dataset: "
        "the general Accuracy Metrics comparing each standalone model's performance, alongside the final Voting Classifier."
    )
    
    if os.path.exists("accuracy_metrics.png"):
        doc.add_picture("accuracy_metrics.png", width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[accuracy_metrics.png not found]")
        
    doc.add_paragraph()
    doc.add_paragraph("Precision Metrics across all 22 crop classification labels:")
    if os.path.exists("precision_matrix.png"):
        doc.add_picture("precision_matrix.png", width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[precision_matrix.png not found]")

    doc.save(output_path)
    print(f"Word document successfully saved to {output_path}")

if __name__ == "__main__":
    flowchart_img = "algorithm_flowchart.png"
    word_doc = "algo_explanation.docx"
    
    print("Generating flowchart...")
    create_flowchart(flowchart_img)
    
    print("Generating Word document...")
    create_word_doc(word_doc, flowchart_img)
